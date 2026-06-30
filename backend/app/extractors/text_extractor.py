import os
import io
import logging
from typing import Union

logger = logging.getLogger(__name__)


def extract_text_from_pdf(content: Union[bytes, str]) -> str:
    """
    Extract raw text from a PDF file (path or bytes) using PyMuPDF (fitz),
    including embedded geometric hyperlinks.
    """
    try:
        import fitz
    except ImportError:
        logger.error("PyMuPDF (fitz) is not installed.")
        return ""

    text_parts = []
    try:
        if isinstance(content, str) and os.path.exists(content):
            doc = fitz.open(content)
        elif isinstance(content, bytes):
            doc = fitz.open(stream=content, filetype="pdf")
        else:
            logger.warning(f"Invalid PDF input type or non-existent path: {type(content)}")
            return ""

        for page in doc:
            page_links = page.get_links()
            
            page_dict = page.get_text("dict")
            page_text_lines = []
            
            for b in page_dict.get("blocks", []):
                if b.get("type") == 0:  # Text block
                    for l in b.get("lines", []):
                        line_text = ""
                        total_chars = 0
                        bold_chars = 0
                        bold_sizes = []
                        
                        # Check if any geometric links intersect this line
                        line_rect = fitz.Rect(l.get("bbox", [0,0,0,0]))
                        line_urls = []
                        for link in page_links:
                            if "from" in link and "uri" in link:
                                if line_rect.intersects(fitz.Rect(link["from"])):
                                    line_urls.append(link["uri"])
                                    
                        for s in l.get("spans", []):
                            span_text = s.get("text", "")
                            if not span_text.strip():
                                line_text += span_text
                                continue
                            flags = s.get("flags", 0)
                            font_name = s.get("font", "").lower()
                            is_bold = bool(flags & 16) or "bold" in font_name
                            line_text += span_text
                            
                            chars = len(span_text.strip())
                            total_chars += chars
                            if is_bold:
                                bold_chars += chars
                                if s.get("size"):
                                    bold_sizes.extend([round(s.get("size"), 1)] * chars)
                                
                        if line_text.strip():
                            # If more than 50% of the line's visible characters are bold, tag the line
                            if total_chars > 0 and (bold_chars / total_chars) >= 0.5:
                                dominant_size = max(set(bold_sizes), key=bold_sizes.count) if bold_sizes else 11.0
                                line_text = f"[BOLD:{dominant_size}] {line_text.strip()}"
                                
                            # Inject links inline
                            if line_urls:
                                unique_urls = list(dict.fromkeys(line_urls))
                                for u in unique_urls:
                                    line_text += f" [LINK:{u}]"
                                    
                            page_text_lines.append(line_text)
                            
            page_text = "\n".join(page_text_lines)
            text_parts.append(page_text)
            
        doc.close()
    except Exception as e:
        logger.error(f"Error reading PDF content: {e}")
        return ""

    return "\n".join(text_parts)


def extract_text_from_docx(content: Union[bytes, str]) -> str:
    """
    Extract raw text from a DOCX file (path or bytes) using python-docx.
    """
    try:
        import docx
    except ImportError:
        logger.error("python-docx is not installed.")
        return ""

    text_parts = []
    try:
        if isinstance(content, str) and os.path.exists(content):
            doc = docx.Document(content)
        elif isinstance(content, bytes):
            doc = docx.Document(io.BytesIO(content))
        else:
            logger.warning(f"Invalid DOCX input type or non-existent path: {type(content)}")
            return ""

        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
    except Exception as e:
        logger.error(f"Error reading DOCX content: {e}")
        return ""

    return "\n".join(text_parts)


def extract_text_from_file(file_path: str, file_content: Union[bytes, None] = None) -> str:
    """
    General utility to extract text given a filename and optional binary contents.
    Dispatches based on extension (.pdf, .docx, .txt).
    """
    ext = os.path.splitext(file_path)[1].lower()

    data = file_content if file_content is not None else file_path

    if ext == ".pdf":
        return extract_text_from_pdf(data)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(data)
    elif ext in (".txt", ".md", ".json", ".csv"):
        if isinstance(data, bytes):
            try:
                return data.decode("utf-8", errors="ignore")
            except Exception as e:
                logger.error(f"Error decoding text file bytes: {e}")
                return ""
        elif isinstance(data, str) and os.path.exists(data):
            try:
                with open(data, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                return ""
        return str(data)
    else:
        logger.warning(f"Unsupported file extension {ext} for text extraction.")
        return ""
